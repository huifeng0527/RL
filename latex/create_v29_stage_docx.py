from pathlib import Path
import re

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from PIL import Image

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "manuscripts" / "methodology_ieee_v2.52_kinematics_formulas_only.md"
OUT_PATH = ROOT / "manuscripts" / "methodology_ieee_v2.29_stage_draft_with_figures_v27.docx"
EQUATION_DIR = ROOT / "manuscripts" / "figures" / "generated_equations_v29"
FIG_DIR = ROOT / "manuscripts" / "figures" / "paper_ready"
FIG_MAP = {
    "Fig. 1.": ROOT / "latex" / "overleaf_rl_rehab_v27_four_figures" / "figures" / "fig_system_framework.png",
    "Fig. 2.": FIG_DIR / "fig2.png",
    "Fig. 3.": FIG_DIR / "fig_sim01_league_overview_no_title.png",
    "Fig. 4.": FIG_DIR / "fig_sim02_ablation_gru_aux_composite_filled_no_title.png",
    "Fig. 5.": FIG_DIR / "fig_physical_deployment.png",
}


def set_run_font(run, font="Times New Roman", size=10):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_md_runs(paragraph, text, base_size=10):
    parts = text.split("**")
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part.replace("\\quad", "    "))
        set_run_font(run, size=base_size)
        run.bold = idx % 2 == 1


def add_md_table(doc, block):
    rows = []
    for raw in block:
        cells = [c.strip() for c in raw.strip().strip("|").split("|")]
        if all(set(c.replace(":", "").replace("-", "").strip()) == set() for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=8.5)
                    run.bold = i == 0
            if i == 0:
                set_cell_shading(cell, "EDEDED")


def add_caption(doc, text):
    p = doc.add_paragraph(style="IEEE Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)


def add_figure(doc, path, key):
    if not path.exists():
        p = doc.add_paragraph(style="IEEE Caption")
        p.add_run(f"[Missing figure: {path.as_posix()}]").bold = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    width = Inches(6.8 if key == "Fig. 2." else 6.6)
    run.add_picture(str(path), width=width)


def clean_equation(expr):
    expr = expr.strip()
    if expr.startswith("$$"):
        expr = expr[2:]
    if expr.endswith("$$"):
        expr = expr[:-2]
    return expr.strip()


def m_run(text):
    run = OxmlElement("m:r")
    text_el = OxmlElement("m:t")
    text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_el.text = text
    run.append(text_el)
    return run


def m_group(items):
    out = []
    for item in items:
        if isinstance(item, str):
            out.append(m_run(item))
        else:
            out.append(item)
    return out


def m_slot(name, items):
    slot = OxmlElement(name)
    for item in m_group(items if isinstance(items, list) else [items]):
        slot.append(item)
    return slot


def m_sub(base, sub):
    el = OxmlElement("m:sSub")
    el.append(m_slot("m:e", base if isinstance(base, list) else [base]))
    el.append(m_slot("m:sub", [sub]))
    return el


def m_sup(base, sup):
    el = OxmlElement("m:sSup")
    el.append(m_slot("m:e", base if isinstance(base, list) else [base]))
    el.append(m_slot("m:sup", [sup]))
    return el


def m_subsup(base, sub, sup):
    el = OxmlElement("m:sSubSup")
    el.append(m_slot("m:e", base if isinstance(base, list) else [base]))
    el.append(m_slot("m:sub", [sub]))
    el.append(m_slot("m:sup", [sup]))
    return el


def m_frac(num, den):
    el = OxmlElement("m:f")
    fpr = OxmlElement("m:fPr")
    ftype = OxmlElement("m:type")
    ftype.set(qn("m:val"), "bar")
    fpr.append(ftype)
    el.append(fpr)
    el.append(m_slot("m:num", num if isinstance(num, list) else [num]))
    el.append(m_slot("m:den", den if isinstance(den, list) else [den]))
    return el


def m_math_para(doc, lines):
    p = doc.add_paragraph(style="Display Equation")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    for line_index, items in enumerate(lines):
        if line_index:
            p.add_run().add_break()
        omath_para = OxmlElement("m:oMathPara")
        omath = OxmlElement("m:oMath")
        for item in m_group(items):
            omath.append(item)
        omath_para.append(omath)
        p._p.append(omath_para)
    return p


def equation_lines(index):
    def diff_terms():
        return [m_sub("p", "R,t"), " − ", m_sub("p", "H,t")]

    def norm_terms():
        return ["‖", *diff_terms(), "‖₂"]

    mapping = {
        1: [
            [m_sub("r", "t"), " = ", m_sub("r", "dist"), "(", m_sub("d", "t"), ") + ", m_sub("r", "smooth"), "(", m_sub("a", "t"), ") + ", m_sub("r", "bound"), "(", m_sub("p", "R,t"), ") + ", m_sub("r", "catch"), "(", m_sub("d", "t"), ")", "    (1)"],
        ],
        2: [
            [m_subsup("u", "t", "SPC"), " = ", m_sub("ρ", "ep"), m_sub("q", "t"), "    (2)"],
        ],
        3: [[m_sub("ṽ", "t"), " = α", m_sub("u", "t"), " + (1−α)", m_sub("v", "t−1"), "    (3)"]],
        4: [[m_sub("v", "t"), " = ", m_sub("v", "t−1"), " + clip_norm(", m_sub("ṽ", "t"), "−", m_sub("v", "t−1"), ", ", m_sub("a", "max"), ")", "    (4)"]],
        5: [[m_sub("p", "H,t+1"), " = ", m_sub("p", "H,t"), " + ", m_sub("v", "t"), "    (5)"]],
        6: [[m_sub("o", "t"), " = [", m_sub("s", "t"), "; ", m_sub("h", "t−T:t"), "]", "    (6)"]],
        7: [[m_sub("s", "t"), " = [", m_sub("p", "R"), "; ", m_sub("p", "H"), "; d(R,H); ", m_sub("b", "N"), ", ", m_sub("b", "S"), ", ", m_sub("b", "E"), ", ", m_sub("b", "W"), "; stride; ", m_sub("a", "t−1"), "]", "    (7)"]],
        8: [[m_sub("L", "traj"), " = E[", m_sup(["‖", m_sub("D̂", "t+1:t+H"), "−", m_sub("D", "t+1:t+H"), "‖₂"], "2"), "],  H=8", "    (8)"]],
        9: [[m_sub("L", "total"), " = ", m_sub("L", "PPO"), " + ", m_sub("λ", "traj"), m_sub("L", "traj"), "    (9)"]],
        10: [[m_sub("s", "i"), " = ", m_frac(["1"], [m_sub("ℓ̄", "i"), " + ", m_sub("ε", "ℓ")]), ",  ", m_sub("p̃", "i"), " = (1−μ)", m_frac([m_sub("s", "i")], ["∑", m_sub("s", "j")]), " + ", m_frac(["μ"], ["m"]), "    (10)"]],
    }
    return mapping.get(index, [[clean_equation("")]])


def add_equation(doc, expr, index):
    m_math_para(doc, equation_lines(index))


def add_algorithm_line(doc, text):
    p = doc.add_paragraph(style="Algorithm Text")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(1)
    add_md_runs(p, text, base_size=9)


def is_algorithm_line(text):
    return (
        text.startswith("**Input:**")
        or text.startswith("**Initialize:**")
        or any(text.startswith(f"{n}.") for n in range(1, 20))
    )


def add_algorithm_block(doc, caption, lines):
    table = doc.add_table(rows=len(lines) + 1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header = table.cell(0, 0)
    set_cell_shading(header, "EDEDED")
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run(caption)
    set_run_font(header_run, size=9)
    header_run.bold = True

    for idx, line in enumerate(lines, start=1):
        cell = table.cell(idx, 0)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if line.startswith(tuple(f"{n}." for n in range(2, 20))):
            p.paragraph_format.left_indent = Inches(0.18)
        add_md_runs(p, line, base_size=9)


def configure_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True

    caption_style = styles.add_style("IEEE Caption", 1)
    caption_style.font.name = "Times New Roman"
    caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption_style.font.size = Pt(9)

    equation_style = styles.add_style("Display Equation", 1)
    equation_style.font.name = "Cambria Math"
    equation_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    equation_style.font.size = Pt(10)

    algorithm_style = styles.add_style("Algorithm Text", 1)
    algorithm_style.font.name = "Times New Roman"
    algorithm_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    algorithm_style.font.size = Pt(9)


def main():
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Stage Manuscript Draft: Methods, Experiments, and Physical Deployment")
    set_run_font(run, size=14)
    run.bold = True

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Abstract, Introduction, Discussion, Conclusion, and References remain to be finalized.")
    set_run_font(r, size=9)
    r.italic = True

    i = 0
    equation_index = 1
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("$$"):
            eq_lines = [stripped]
            if not (stripped.endswith("$$") and len(stripped) > 2):
                i += 1
                while i < len(lines):
                    eq_lines.append(lines[i].strip())
                    if lines[i].strip().endswith("$$"):
                        break
                    i += 1
            add_equation(doc, " ".join(eq_lines), equation_index)
            equation_index += 1
            i += 1
            continue

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_md_table(doc, block)
            continue

        if stripped.startswith("# "):
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue

        if stripped.startswith("**Fig.") and stripped.endswith("**"):
            caption = stripped.strip("*")
            for key, fig_path in FIG_MAP.items():
                if caption.startswith(key):
                    add_figure(doc, fig_path, key)
                    break
            add_caption(doc, caption)
            i += 1
            continue

        if stripped.startswith("**TABLE") and stripped.endswith("**"):
            p = doc.add_paragraph(style="IEEE Caption")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped.strip("*"))
            set_run_font(run, size=9)
            run.bold = True
            i += 1
            continue

        if stripped.startswith("**Algorithm 1:"):
            caption = stripped.strip("*")
            i += 1
            algorithm_lines = []
            while i < len(lines):
                candidate = lines[i].strip()
                if not candidate:
                    i += 1
                    continue
                if is_algorithm_line(candidate):
                    algorithm_lines.append(candidate)
                    i += 1
                    continue
                break
            add_algorithm_block(doc, caption, algorithm_lines)
            continue

        if is_algorithm_line(stripped):
            add_algorithm_line(doc, stripped)
            i += 1
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.2)
        add_md_runs(p, stripped)
        i += 1

    doc.save(OUT_PATH)
    print(OUT_PATH.as_posix())


if __name__ == "__main__":
    main()
