from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = Path('manuscripts/methodology_ieee_v2.6_revised_simulation.docx')
out = Path('manuscripts/methodology_ieee_v2.6_revised_simulation.docx')

doc = Document(path)

paragraph_replacements = {
    'We evaluate the simulation component through two complementary studies. First, a ten-generation league experiment tests whether iterative opponent-pool training improves robustness against heterogeneous learned hand behaviors. Second, a GRU-based network ablation isolates the effect of temporal interaction-history encoding and the auxiliary future-dynamics head. These experiments replace the earlier LSTM ablation and OOD tables in the previous manuscript version, because the current paper scope centers on the no-opponent-ID league result and the final GRU ablation trained against H1-H10 plus scripted hands.':
    'We evaluate the simulation component through three complementary studies. First, a ten-generation league experiment examines whether iterative opponent-pool training improves robustness against heterogeneous learned hand behaviors. Second, a GRU-based network ablation isolates the contribution of temporal interaction-history encoding and auxiliary future-dynamics supervision. Third, a held-out policy comparison evaluates the resulting robot policies against both a scripted pursuit controller and a learned hand agent.',

    'The architecture ablation compares three robot policy encoders trained for one generation against the same H1-H10 plus scripted-hand pool. MLP uses only the scalar spatial state and ignores interaction history. GRU adds a temporal sequence encoder over the 16-frame relative-displacement buffer. GRU+Aux further adds the auxiliary future-prediction objective described in Section IV-B. Each group is trained for 8M environment steps. Raw per-episode reward, episode length, ZPD steps, and workspace coverage are recorded independently of TensorBoard.':
    'The architecture ablation compares three robot policy encoders trained for one generation under the same mixed-opponent protocol. The MLP baseline uses only the scalar spatial state and does not process the interaction-history buffer. The GRU variant adds a temporal sequence encoder over the 16-frame relative-displacement history. The GRU+Aux variant further augments the recurrent encoder with the auxiliary future-prediction objective described in Section IV-B. All variants are trained for 8 million environment steps against a pool composed of the first ten learned hand policies and the scripted pursuit controller.',

    'Fig. 2 shows the updated ablation result. The top row uses raw per-episode training data smoothed over a 200k-timestep window. The MLP baseline plateaus at lower episode reward and shorter episode length than both recurrent variants. In the final 10% of training episodes, MLP obtains a mean reward of -42.94 and mean episode length of 27.58, whereas GRU reaches -35.56 reward and 36.12 episode length. This corresponds to an improvement of 7.39 reward points and 8.54 steps per episode. GRU+Aux reaches a similar final-stage reward (-35.39) and episode length (35.92), indicating that temporal sequence modeling is the dominant contributor to aggregate control performance in this run.':
    'Fig. 2 summarizes the GRU ablation. The top row reports smoothed training trajectories for episode return and episode length, using a 200k-timestep smoothing window to emphasize the long-term learning trend. The MLP baseline converges to lower return and shorter episodes than both recurrent variants, indicating that instantaneous geometric state alone is insufficient for stable closed-loop adaptation. Over the final 10% of training episodes, the MLP obtains a mean reward of -42.94 and a mean episode length of 27.58, whereas the GRU reaches -35.56 reward and 36.12 steps. This corresponds to an improvement of 7.39 reward points and 8.54 steps per episode. The GRU+Aux policy reaches a comparable final-stage reward (-35.39) and episode length (35.92), suggesting that temporal sequence modeling is the primary source of aggregate control improvement in this experiment.',

    'The auxiliary model does not clearly dominate GRU-only on final aggregate metrics, but it provides an interpretable representation-learning signal. The bottom row of Fig. 2 visualizes future-motion prediction examples and the trajectory error over the eight-step horizon. The error increases with prediction horizon, as expected for multi-step forecasting, while early-step predictions remain close to the observed hand trajectory. These results support the auxiliary task as a mechanism for encouraging the encoder to internalize latent patient dynamics, even though additional tuning may be required for a larger control-performance gain.':
    'Although the auxiliary model does not substantially exceed the GRU-only policy in final aggregate control metrics, it provides a direct and interpretable representation-learning signal. The bottom row of Fig. 2 visualizes representative future-motion predictions and the trajectory error over the eight-step horizon. The error increases with prediction horizon, as expected for multi-step forecasting, while short-horizon predictions remain close to the observed hand trajectory. These results indicate that the auxiliary task encourages the encoder to represent latent hand dynamics, even when the resulting control-performance gain is modest under the current loss weighting.',

    'Fig. 2. GRU ablation and auxiliary future-motion prediction. Panels (a) and (b) show smoothed raw training reward and episode length for MLP, GRU, and GRU+Aux policies trained against the H1-H10 plus scripted-hand pool. Panels (c) and (d) visualize representative auxiliary future-prediction examples and horizon-dependent trajectory error.':
    'Fig. 2. GRU ablation and auxiliary future-motion prediction. Panels (a) and (b) show smoothed training reward and episode length for MLP, GRU, and GRU+Aux policies trained against the mixed learned-hand and scripted-hand pool. Panels (c) and (d) visualize representative auxiliary future-prediction examples and horizon-dependent trajectory error.',

    'The revised simulation results support two main conclusions. First, league training improves robustness by exposing the robot to a distribution of learned patient behaviors rather than a single scripted controller. The improvement in both mean and worst-case TIS suggests that opponent-pool training reduces exploitability and better approximates patient variability. Second, temporal interaction history is essential for adaptive control in this task. The gap between MLP and GRU policies shows that instantaneous spatial state alone is insufficient for inferring patient motion trends and maintaining sustained interaction.':
    'The revised simulation results support three main conclusions. First, league training improves robustness by exposing the robot to a distribution of learned patient behaviors rather than a single scripted controller. The improvement in both mean and worst-case TIS suggests that opponent-pool training reduces exploitability and better approximates patient variability. Second, temporal interaction history is essential for adaptive control in this task. The gap between MLP and GRU policies shows that instantaneous spatial state alone is insufficient for inferring patient motion trends and maintaining sustained interaction. Third, evaluation against both scripted and learned-agent hands reveals complementary failure modes: scripted-only training performs reasonably against the scripted controller but generalizes poorly to the learned agent, while single-agent training over-specializes to the learned hand and degrades on the scripted controller. The league-trained policy provides the most balanced behavior across the two test conditions.',
}

for p in doc.paragraphs:
    text = p.text.strip()
    if text in paragraph_replacements:
        p.text = paragraph_replacements[text]

# Remove the GRU ablation summary table heading, table, and explanatory paragraph.
body = doc._body._element
children = list(body)
remove_indices = []
for idx, child in enumerate(children):
    if child.tag.endswith('p'):
        text = ''.join(node.text for node in child.iter() if node.text).strip()
        if text == 'TABLE II: FINAL-STAGE GRU ABLATION SUMMARY':
            remove_indices.append(idx)
            if idx + 1 < len(children) and children[idx + 1].tag.endswith('tbl'):
                remove_indices.append(idx + 1)
            if idx + 2 < len(children) and children[idx + 2].tag.endswith('p'):
                next_text = ''.join(node.text for node in children[idx + 2].iter() if node.text).strip()
                if next_text.startswith('Table II reports'):
                    remove_indices.append(idx + 2)
            break
for idx in sorted(set(remove_indices), reverse=True):
    body.remove(children[idx])

# Insert the policy-comparison section before Sim-to-Real Transfer.
def insert_paragraph_before(paragraph, text, bold=False):
    new_p = paragraph._element.addprevious(deepcopy(paragraph._element))
    # addprevious returns None in lxml; retrieve previous sibling
    new_el = paragraph._element.getprevious()
    new_para = paragraph._parent.paragraphs[0]
    for p in paragraph._parent.paragraphs:
        if p._element is new_el:
            new_para = p
            break
    new_para.clear()
    run = new_para.add_run(text)
    run.bold = bold
    return new_para

def insert_table_before(paragraph, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = 'Table Grid'
    for cell, value in zip(table.rows[0].cells, rows[0]):
        cell.text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    paragraph._element.addprevious(table._element)
    return table

sim_to_real = None
for p in doc.paragraphs:
    if p.text.strip() == 'C. Sim-to-Real Transfer':
        sim_to_real = p
        break
if sim_to_real is None:
    raise RuntimeError('Could not locate Sim-to-Real section')

sim_to_real.text = 'D. Sim-to-Real Transfer'

# Insert in reverse order because each insert is before the same anchor.
insert_paragraph_before(sim_to_real, 'The learned-agent test highlights the main benefit of league training. The scripted-only baseline attains a TIS of 0.19 and a 90% catch rate against the learned hand, indicating poor robustness outside its training distribution. The single-agent baseline performs well against the learned hand (TIS = 0.49) but drops sharply on the scripted hand (TIS = 0.11), suggesting over-specialization to one learned opponent. The league-trained robot achieves competitive TIS on the learned hand (0.48) while also retaining the strongest scripted-hand TIS (0.38), yielding the most balanced performance across the two non-interactive tests. Mouse-controlled human-in-the-loop evaluation will be reported separately after collection.')
insert_table_before(sim_to_real, [
    ['Robot policy', 'Scripted hand', 'Agent hand'],
    ['Scripted-only', '0.31 / 0.54 / 52 / 60%', '0.19 / 0.54 / 31 / 90%'],
    ['Single-agent', '0.11 / 0.41 / 25 / 95%', '0.49 / 0.60 / 70 / 38%'],
    ['League', '0.38 / 0.55 / 67 / 43%', '0.48 / 0.65 / 64 / 48%'],
])
insert_paragraph_before(sim_to_real, 'TABLE II: POLICY COMPARISON ON SCRIPTED AND LEARNED-AGENT HANDS')
insert_paragraph_before(sim_to_real, 'Table II reports the final policy comparison using the format TIS / ZPD coverage / episode length / catch rate. The scripted-hand test uses the environment\'s stochastic pursuit controller, whereas the agent-hand test uses a learned hand policy from the league pool.')
insert_paragraph_before(sim_to_real, 'C. Policy Comparison on Scripted and Learned-Agent Hands', bold=True)

# Rename conclusion wording to avoid overclaiming and match current results.
for p in doc.paragraphs:
    if p.text.strip() == 'This manuscript presents a simulation-to-real RL framework for adaptive non-contact upper-limb rehabilitation. CMD-DR generates diverse virtual patient behavior by separating intended movement strategy from motor execution constraints. A dual-stream recurrent encoder uses interaction history to infer latent patient dynamics, and an auxiliary future-prediction head supplies dense self-supervised representation learning. Iterative league training improves robustness across learned hand opponents, and the GRU ablation confirms the importance of temporal sequence modeling. The next stage is to integrate these simulation results with systematic physical validation on the UR10 magnetic microrobot platform.':
        p.text = 'This manuscript presents a simulation-to-real RL framework for adaptive non-contact upper-limb rehabilitation. CMD-DR generates diverse virtual patient behavior by separating intended movement strategy from motor execution constraints. A dual-stream recurrent encoder uses interaction history to infer latent patient dynamics, and an auxiliary future-prediction head supplies dense self-supervised representation learning. Iterative league training improves robustness across learned hand opponents, and the GRU ablation confirms the importance of temporal sequence modeling. The scripted-versus-agent comparison further shows that league training provides more balanced behavior than policies trained against a single opponent type. The next stage is to complete mouse-controlled testing and systematic physical validation on the UR10 magnetic microrobot platform.'

doc.save(out)
print(out.as_posix())
