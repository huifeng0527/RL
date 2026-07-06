from pathlib import Path
from shutil import copy2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

src = Path("manuscripts/methodology_ieee_v2.10_experiment_flow.docx")
out = Path("manuscripts/methodology_ieee_v2.11_league_clarity.docx")
league_fig = Path("manuscripts/figures/paper_ready/fig_sim01_league_overview_no_title.png")
ablation_fig = Path("manuscripts/figures/paper_ready/fig_sim02_ablation_gru_aux_composite_filled_no_title.png")
copy2(src, out)
doc = Document(out)

body = doc._body._element
children = list(body)
start_idx = None
for idx, child in enumerate(children):
    if child.tag.endswith("p"):
        text = "".join(node.text for node in child.iter() if node.text).strip()
        if text.startswith("VI. EXPERIMENTS"):
            start_idx = idx
            break
if start_idx is None:
    raise RuntimeError("Could not find Section VI")
for child in children[start_idx:]:
    if child.tag.endswith("sectPr"):
        continue
    body.remove(child)


def add_heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


def add_para(text):
    return doc.add_paragraph(text)


def add_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
    return p


def add_figure(path, caption, width=6.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(caption)


add_heading("VI. EXPERIMENTS AND RESULTS")
add_para("The experiments evaluate the proposed framework from complementary perspectives rather than following the methodology section mechanically. We first define the common simulation setup, training protocols, and evaluation metrics. The main evaluation then examines whether league training improves robustness across learned hand behaviors and whether this robustness transfers across scripted and learned-agent hand controllers. We next analyze the role of temporal interaction history and the auxiliary future-dynamics head through a network ablation. The section closes with the real-world deployment pipeline that connects the simulation policy to the UR10 platform.")

add_heading("A. Experimental Setup")
add_para("All simulation experiments use the pursuit-evasion rehabilitation environment described in Section III. The robot controls the magnetic microrobot, and the hand is controlled either by a scripted pursuit controller or by a learned hand policy. The therapeutic objective is defined by a ZPD distance band of 3.5-5.5 workspace units. A step is counted as in-zone when the hand-microrobot distance lies inside this band.")
add_para("The primary evaluation metric is Time-In-Zone (TIS), the fraction of the episode horizon during which the interaction remains inside the ZPD band. We also report ZPD coverage, episode length, catch rate, and empirical robustness across the learned hand pool. These metrics capture different aspects of the same rehabilitation objective: TIS measures sustained therapeutic interaction over the full horizon, ZPD coverage measures the quality of the realized portion of an episode, episode length captures survival, and catch rate reflects overly easy interactions in which the microrobot is caught too quickly.")
add_para("We compare three representative robot training protocols. The scripted-only baseline is trained against the stochastic scripted pursuit controller. The single-agent baseline is trained against one learned hand policy. The league policy is trained through iterative opponent-pool expansion, where robot generations are trained against a mixture of scripted and learned hand behaviors. For the network ablation, we compare an MLP policy, a GRU policy with temporal interaction history, and a GRU+Aux policy with the auxiliary future-dynamics head.")

add_heading("B. League Training and Robustness Evaluation")
add_para("This evaluation tests whether opponent-pool training reduces the brittle specialization that can occur when the robot is optimized against a single hand model. The league policy is trained for ten generations while the hand pool is expanded with learned hand policies. During robot training, the opponent identity is not given as an input; the robot must infer the current hand behavior from the spatial state and recent displacement history. This design makes the evaluation closer to rehabilitation use, where patient capability is observed through movement rather than provided as a label.")
add_para("The league analysis in Fig. 1 is computed by freezing each robot generation and evaluating it against the learned hand pool under the same ZPD band. For each robot-hand pairing, we measure the fraction of time spent inside the ZPD and then summarize performance in three ways: the mean trend over the pool, the lower-tail or worst-case trend, and an empirical matrix of pairwise outcomes. The mean curve indicates overall adaptation, whereas the worst-case curve is used to detect whether some hand policies remain failure cases. The matrix provides the most detailed view because it shows which robot generations are robust across easy and difficult virtual patients.")
add_para("The resulting pattern is not only an increase in average performance. Later league generations also improve the lower tail, meaning that the robot becomes less dependent on favorable hand behaviors. This distinction is important for rehabilitation: a policy that works well for typical patients but collapses for difficult movement patterns would not provide reliable personalized training. The sorted matrix supports the same interpretation, showing that later generations maintain useful ZPD interaction across a broader portion of the learned-hand pool.")
add_para("The scripted-hand curve in Fig. 1 is used as a consistency check rather than the main robustness test. The scripted controller follows a noisy pursuit rule and therefore probes whether league training preserves basic tracking behavior. The learned-hand matrix is more diagnostic because learned opponents can expose strategy-dependent weaknesses. Taken together, the figure indicates that opponent accumulation improves robustness without relying on an explicit opponent label.")
add_figure(league_fig, "Fig. 1. League-training overview for the no-opponent-ID ZPD 3.5-5.5 simulation. The final robot generation improves mean and worst-case Time-In-Zone against the learned-hand pool, while the sorted empirical matrix visualizes robustness across increasingly difficult virtual patients.")
add_para("We further test whether this robustness transfers across hand-controller types. For this comparison, the final scripted-only, single-agent, and league-trained robot policies are each evaluated against two hand mechanisms: the stochastic scripted pursuit controller and a learned hand agent. The table entries report TIS, ZPD coverage, episode length, and catch rate in that order. Mouse-controlled human-in-the-loop testing is kept separate because it introduces human reaction time and voluntary strategy rather than another automated controller.")
add_heading("TABLE II: POLICY COMPARISON ON SCRIPTED AND LEARNED-AGENT HANDS")
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
for cell, value in zip(table.rows[0].cells, ["Robot policy", "Scripted hand", "Agent hand"]):
    cell.text = value
for row in [
    ["Scripted-only", "0.31 / 0.54 / 52 / 60%", "0.19 / 0.54 / 31 / 90%"],
    ["Single-agent", "0.11 / 0.41 / 25 / 95%", "0.49 / 0.60 / 70 / 38%"],
    ["League", "0.38 / 0.55 / 67 / 43%", "0.48 / 0.65 / 64 / 48%"],
]:
    cells = table.add_row().cells
    for cell, value in zip(cells, row):
        cell.text = value
add_para("Table II shows complementary failure modes for the two non-league baselines. The scripted-only robot remains competitive on the scripted controller but is caught frequently by the learned hand agent, indicating that rule-based pursuit during training does not cover learned strategic behavior. The single-agent robot shows the reverse tendency: it performs well against the learned hand but degrades on the scripted controller. The league-trained robot does not dominate every metric in every column, but it avoids the severe collapse observed in the baselines and provides the most balanced behavior across both automated test conditions. This balance is the relevant outcome for an adaptive rehabilitation controller, whose goal is not to overfit to one hand model but to maintain therapeutic interaction as patient behavior changes.")

add_heading("C. Network Ablation and Auxiliary Dynamics Analysis")
add_para("The ablation study turns from opponent diversity to policy representation. The key question is whether the robot can regulate difficulty from the current geometric state alone, or whether it needs recent patient-motion history. The MLP baseline observes the current geometric state but does not explicitly encode recent interaction history. The GRU policy processes the 16-frame relative-displacement buffer, allowing it to infer hand velocity, response delay, and pursuit tendency. The GRU+Aux policy further adds the future-dynamics prediction head described in Section IV-B.")
add_para("The upper panels of Fig. 2 show that the main performance gain appears when temporal interaction history is introduced. Both recurrent policies learn longer and more rewarding interactions than the MLP baseline. This suggests that the robot benefits from observing how the hand has been moving, not only where it is at the current instant. In this setting, the GRU-only and GRU+Aux policies follow similar aggregate learning trends, indicating that temporal sequence modeling is the dominant contributor to control performance.")
add_para("The lower panels of Fig. 2 provide a complementary interpretation of the auxiliary head. Rather than treating the auxiliary task as a separate source of reward improvement, we use it to inspect whether the temporal encoder has learned predictive structure in hand motion. The predicted trajectories capture short-horizon movement direction and diverge gradually at longer horizons, consistent with the uncertainty of multi-step prediction in an interactive task. Thus, the auxiliary head mainly serves to encourage and expose a temporal representation of patient dynamics.")
add_figure(ablation_fig, "Fig. 2. GRU ablation and auxiliary future-motion prediction. Panels (a) and (b) show smoothed training reward and episode length for MLP, GRU, and GRU+Aux policies trained against the mixed learned-hand and scripted-hand pool. Panels (c) and (d) visualize representative auxiliary future-prediction examples and horizon-dependent trajectory error.")

add_heading("D. Real-World Deployment Pipeline")
add_para("The final part of the experiment section describes how the simulation policy is connected to the physical UR10 platform. This pipeline is presented as an implementation bridge rather than as a paired sim-to-real benchmark, because real hand kinematics, camera latency, magnetic actuation, and safety constraints are not fully represented in the virtual training environment.")
add_para("The physical implementation uses the same 44-dimensional observation structure as simulation: robot position from Real-Time Data Exchange (RTDE), hand and microrobot positions from the overhead camera, boundary features, previous action, and the displacement-history buffer. Perception and control run in separate threads to avoid jitter. The camera runs at 10-15 Hz, whereas the robot requires deterministic commands at a fixed control rate. If both processes ran in one thread, a slow You Only Look Once (YOLO) inference step could stall the control loop.")
add_para("The system therefore uses two threads connected by a single-element message queue. The vision thread captures frames, runs undistortion and detection, and pushes results to the queue. If the queue is full, stale results are replaced. The control thread runs at 20 Hz and attempts a non-blocking read each cycle. If a new frame is available, the hand position is updated; otherwise, the system uses dead reckoning based on a low-pass filtered velocity estimate from recent vision updates. The resulting observation is passed to the PPO policy, and the two-dimensional action is mapped through the inverse homography before being sent to the UR10 through the servoL interface.")


def set_run_font(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        rfonts.set(ns + key, "Times New Roman")


for style in doc.styles:
    if hasattr(style, "font"):
        style.font.name = "Times New Roman"
        style.font.size = Pt(10)

for p in doc.paragraphs:
    for run in p.runs:
        set_run_font(run)

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run)

for section in doc.sections:
    for part in [section.header, section.footer]:
        for p in part.paragraphs:
            for run in p.runs:
                set_run_font(run)

doc.save(out)
print(out.as_posix())
