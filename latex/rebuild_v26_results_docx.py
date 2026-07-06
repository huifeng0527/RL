from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = Path('manuscripts/methodology_ieee_v2.6_revised_simulation.docx')
league_fig = Path('manuscripts/figures/paper_ready/fig_sim01_league_overview_no_title.png')
ablation_fig = Path('manuscripts/figures/paper_ready/fig_sim02_ablation_gru_aux_composite_filled_no_title.png')

doc = Document(path)
body = doc._body._element
children = list(body)

start_idx = None
for idx, child in enumerate(children):
    if child.tag.endswith('p'):
        text = ''.join(node.text for node in child.iter() if node.text).strip()
        if 'VI. EXPERIMENTS AND RESULTS' in text:
            start_idx = idx
            break
if start_idx is None:
    raise RuntimeError('Could not find VI. EXPERIMENTS AND RESULTS')

for child in children[start_idx:]:
    if child.tag.endswith('sectPr'):
        continue
    body.remove(child)


def add_heading(text):
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    return p


def add_para(text):
    return doc.add_paragraph(text)


def add_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
    return p


def add_figure(path, caption, width=6.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(caption)


add_heading('VI. EXPERIMENTS AND RESULTS')
add_para('We evaluate the simulation component through three complementary studies. First, a ten-generation league experiment examines whether iterative opponent-pool training improves robustness against heterogeneous learned hand behaviors. Second, a GRU-based network ablation isolates the contribution of temporal interaction-history encoding and auxiliary future-dynamics supervision. Third, a held-out policy comparison evaluates the resulting robot policies against both a scripted pursuit controller and a learned hand agent.')

add_heading('A. League Training Robustness')
add_para('The league experiment uses a ZPD band of 3.5-5.5 workspace units and trains ten robot generations without providing an explicit opponent identity. Each robot generation is evaluated against the first ten learned hand policies and a scripted pursuit controller. Time-in-zone (TIS), defined as the fraction of interaction time spent inside the ZPD band, is used as the main robustness metric. We report both mean TIS across learned hands and worst-case TIS to measure the lower tail of performance.')
add_para('Fig. 1 summarizes the final ten-generation league result. The final robot R10 achieves the strongest learned-hand robustness among the evaluated generations. Mean TIS against learned hands increases from 0.188 for R1 to 0.282 for R10, an absolute gain of 0.094. Worst-case learned-hand TIS increases from 0.138 to 0.204, an absolute gain of 0.065. These gains indicate that iterative opponent accumulation improves not only average performance but also robustness against harder learned hand policies.')
add_para('The sorted empirical matrix further supports this interpretation. When robots are ordered from weak to strong and hands from easy to hard, later robot generations occupy higher-TIS regions across a broader set of opponents. The scripted-hand curve remains useful as a sanity check, but the learned-hand pool exposes robustness differences that a single scripted hand would not reveal. This result motivates league-trained virtual patients as a stress-test distribution for adaptive rehabilitation policies.')
add_figure(league_fig, 'Fig. 1. League-training overview for the no-opponent-ID ZPD 3.5-5.5 simulation. The final robot generation improves mean and worst-case time-in-zone against the learned-hand pool, while the sorted empirical matrix visualizes robustness across increasingly difficult virtual patients.')

add_heading('B. GRU Network Ablation')
add_para('The architecture ablation compares three robot policy encoders trained for one generation under the same mixed-opponent protocol. The MLP baseline uses only the scalar spatial state and does not process the interaction-history buffer. The GRU variant adds a temporal sequence encoder over the 16-frame relative-displacement history. The GRU+Aux variant further augments the recurrent encoder with the auxiliary future-prediction objective described in Section IV-B. All variants are trained for 8 million environment steps against a pool composed of the first ten learned hand policies and the scripted pursuit controller.')
add_para('Fig. 2 summarizes the GRU ablation. The top row reports smoothed training trajectories for episode return and episode length, using a 200k-timestep smoothing window to emphasize the long-term learning trend. The MLP baseline converges to lower return and shorter episodes than both recurrent variants, indicating that instantaneous geometric state alone is insufficient for stable closed-loop adaptation. Over the final 10% of training episodes, the MLP obtains a mean reward of -42.94 and a mean episode length of 27.58, whereas the GRU reaches -35.56 reward and 36.12 steps. This corresponds to an improvement of 7.39 reward points and 8.54 steps per episode. The GRU+Aux policy reaches a comparable final-stage reward (-35.39) and episode length (35.92), suggesting that temporal sequence modeling is the primary source of aggregate control improvement in this experiment.')
add_para('Although the auxiliary model does not substantially exceed the GRU-only policy in final aggregate control metrics, it provides a direct and interpretable representation-learning signal. The bottom row of Fig. 2 visualizes representative future-motion predictions and the trajectory error over the eight-step horizon. The error increases with prediction horizon, as expected for multi-step forecasting, while short-horizon predictions remain close to the observed hand trajectory. These results indicate that the auxiliary task encourages the encoder to represent latent hand dynamics, even when the resulting control-performance gain is modest under the current loss weighting.')
add_figure(ablation_fig, 'Fig. 2. GRU ablation and auxiliary future-motion prediction. Panels (a) and (b) show smoothed training reward and episode length for MLP, GRU, and GRU+Aux policies trained against the mixed learned-hand and scripted-hand pool. Panels (c) and (d) visualize representative auxiliary future-prediction examples and horizon-dependent trajectory error.')

add_heading('C. Policy Comparison on Scripted and Learned-Agent Hands')
add_para('To complement the training curves, we evaluate the final robot policies against two representative non-interactive test conditions: a stochastic scripted pursuit controller and a learned hand agent. This comparison probes whether each training protocol produces behavior that is specific to its training opponent or remains effective across different hand-control mechanisms.')
add_heading('TABLE II: POLICY COMPARISON ON SCRIPTED AND LEARNED-AGENT HANDS')
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
for cell, value in zip(table.rows[0].cells, ['Robot policy', 'Scripted hand', 'Agent hand']):
    cell.text = value
for row in [
    ['Scripted-only', '0.31 / 0.54 / 52 / 60%', '0.19 / 0.54 / 31 / 90%'],
    ['Single-agent', '0.11 / 0.41 / 25 / 95%', '0.49 / 0.60 / 70 / 38%'],
    ['League', '0.38 / 0.55 / 67 / 43%', '0.48 / 0.65 / 64 / 48%'],
]:
    cells = table.add_row().cells
    for cell, value in zip(cells, row):
        cell.text = value
add_para('The table reports TIS / ZPD coverage / episode length / catch rate. The scripted-only baseline attains a TIS of 0.31 against the scripted hand but drops to 0.19 TIS with a 90% catch rate against the learned hand, indicating weak robustness outside its training distribution. The single-agent baseline performs well against the learned hand (TIS = 0.49) but degrades substantially on the scripted hand (TIS = 0.11), suggesting over-specialization to one learned opponent. The league-trained robot achieves competitive TIS on the learned hand (0.48) while also retaining the strongest scripted-hand TIS (0.38), yielding the most balanced performance across the two automated test conditions. Mouse-controlled human-in-the-loop evaluation will be reported after collection.')

add_heading('D. Sim-to-Real Transfer')
add_para('The physical UR10 deployment is not directly comparable to simulation through a pointwise sim-real gap table, because real human behavior, camera latency, magnetic actuation, and safety constraints are not distribution-matched to virtual hands. We therefore treat sim-to-real validation as a real-world closed-loop feasibility study rather than as a paired simulation benchmark. The implementation uses the same 44-dimensional observation structure as simulation: robot position from RTDE, hand and microrobot positions from the overhead camera, boundary features, previous action, and the displacement-history buffer.')
add_para('The implementation uses temporal decoupling between perception and control. The camera runs at 10-15 Hz, but the robot needs deterministic commands at a fixed control rate. Running both in one thread would cause jitter: a slow YOLO inference would stall the control loop and freeze the robot. We therefore use two threads connected by a single-element message queue. The vision thread captures frames, runs undistortion and detection, and pushes results to the queue. If the queue is full, the stale result is replaced. The control thread runs at 20 Hz and attempts a non-blocking read each cycle. If a new frame is available, the hand position is updated; otherwise the system falls back to dead reckoning using a low-pass filtered velocity estimate from recent vision updates.')
add_para('The 44-dimensional observation is assembled from the estimated hand position, the robot position via RTDE, and the displacement-history buffer, then fed to the PPO policy. The 2D action is accumulated as a virtual target in pixel space, rate-limited to prevent large per-step motion, and mapped to world coordinates via the inverse homography. The final pose is sent to the UR10 through the servoL interface with a 0.1 s lookahead time and proportional gain of 400, yielding smooth end-effector tracking.')

add_heading('VII. DISCUSSION')
add_para('The revised simulation results support three main conclusions. First, league training improves robustness by exposing the robot to a distribution of learned patient behaviors rather than a single scripted controller. The improvement in both mean and worst-case TIS suggests that opponent-pool training reduces exploitability and better approximates patient variability. Second, temporal interaction history is essential for adaptive control in this task. The gap between MLP and GRU policies shows that instantaneous spatial state alone is insufficient for inferring patient motion trends and maintaining sustained interaction. Third, evaluation against both scripted and learned-agent hands reveals complementary failure modes: scripted-only training performs reasonably against the scripted controller but generalizes poorly to the learned agent, while single-agent training over-specializes to the learned hand and degrades on the scripted controller. The league-trained policy provides the most balanced behavior across the two test conditions.')
add_para('The auxiliary future-prediction task provides a mechanistic and visually interpretable way to verify that the temporal encoder is learning patient-motion dynamics. However, the current GRU+Aux run does not substantially exceed the GRU-only model in final reward, episode length, or ZPD steps. This indicates that auxiliary supervision is useful but not yet fully optimized. Future work should tune the auxiliary loss weights, prediction horizon, and training curriculum to better convert representation quality into control-performance improvement.')

add_heading('VIII. CONCLUSION')
add_para('This manuscript presents a simulation-to-real RL framework for adaptive non-contact upper-limb rehabilitation. CMD-DR generates diverse virtual patient behavior by separating intended movement strategy from motor execution constraints. A dual-stream recurrent encoder uses interaction history to infer latent patient dynamics, and an auxiliary future-prediction head supplies dense self-supervised representation learning. Iterative league training improves robustness across learned hand opponents, and the GRU ablation confirms the importance of temporal sequence modeling. The scripted-versus-agent comparison further shows that league training provides more balanced behavior than policies trained against a single opponent type. The next stage is to complete mouse-controlled testing and systematic physical validation on the UR10 magnetic microrobot platform.')

doc.save(path)
print(path.as_posix())
