from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

src = Path('manuscripts/methodology_ieee_v2.6.docx')
out = Path('manuscripts/methodology_ieee_v2.6_revised_simulation.docx')
league_fig = Path('manuscripts/figures/paper_ready/fig_sim01_league_overview_no_title.png')
ablation_fig = Path('manuscripts/figures/paper_ready/fig_sim02_ablation_gru_aux_composite_filled_no_title.png')

doc = Document(src)

replacements = {
    'The temporal buffer h_{t-T:t} contains the T = 16 most recent relative hand displacement vectors, forming a 32-dimensional sequence. This buffer captures recent velocity, acceleration, and repeated movement patterns that cannot be obtained from the current position alone. The scalar vector is processed by an MLP to extract the current geometric information. In parallel, the temporal buffer is processed by a two-layer LSTM. Since the LSTM uses relative displacements rather than absolute positions, it focuses on movement direction, speed, acceleration, and oscillation. The outputs of the two streams are then concatenated and passed through a fusion MLP to produce a shared representation for policy and value prediction.':
    'The temporal buffer h_{t-T:t} contains the T = 16 most recent relative hand displacement vectors, forming a 32-dimensional sequence. This buffer captures recent velocity, acceleration, and repeated movement patterns that cannot be obtained from the current position alone. The scalar vector is processed by an MLP to extract the current geometric information. In parallel, the temporal buffer is processed by a GRU sequence encoder. Because the GRU uses relative displacements rather than absolute positions, it focuses on movement direction, speed, acceleration, and oscillation. The outputs of the two streams are then concatenated and passed through a fusion MLP to produce a shared representation for policy and value prediction.',
    "Although the ZPD-based reward provides a continuous training signal, it does not directly teach the encoder how to represent patient motion. Learning this representation only through policy reward can be slow and unstable. To provide a more direct signal, we attach a forward dynamics head to the fused representation. This head predicts the patient's next-frame relative displacement. Since the true displacement is available at the next timestep, this target can be obtained without extra data collection or manual annotation. The auxiliary loss is defined as:":
    "Although the ZPD-based reward provides a continuous training signal, it does not directly teach the encoder how to represent patient motion. Learning this representation only through policy reward can be slow and unstable. To provide a more direct signal, we attach a future-dynamics head to the fused representation. This head predicts the patient's future relative displacement over an eight-step horizon and also estimates near-catch risk. Since future displacements are already contained in the rollout, these targets can be obtained without extra data collection or manual annotation. The auxiliary trajectory loss is defined as:",
    'L_aux = E[||d_hat_{t+1} - d_{t+1}||_2^2]':
    'L_traj = E[||D_hat_{t+1:t+H} - D_{t+1:t+H}||_2^2],   H = 8',
    'where d_hat_{t+1} is the predicted next displacement and d_{t+1} is the observed displacement. The full training objective combines the PPO loss with the value loss, entropy regularization, and the auxiliary prediction loss:':
    'where D_hat_{t+1:t+H} is the predicted future displacement sequence and D_{t+1:t+H} is the observed future displacement sequence. The full training objective combines the PPO loss with the auxiliary trajectory and risk-prediction losses:',
    'L_total = L_policy + c_v L_value + c_e L_entropy + lambda L_aux':
    'L_total = L_PPO + lambda_traj L_traj + lambda_risk L_risk',
    'where lambda controls the weight of the auxiliary loss. Because the forward dynamics head shares the encoder with the policy, its gradients also update the spatial and temporal streams. This encourages the encoder to capture patient inertia, delay, and interaction dynamics, helping the robot policy anticipate near-future hand motion instead of reacting only to the current hand-microrobot distance.':
    'where lambda_traj and lambda_risk control the auxiliary loss weights. Because the future-dynamics head shares the encoder with the policy, its gradients also update the spatial and temporal streams. This encourages the encoder to capture patient inertia, delay, and interaction dynamics, helping the robot policy anticipate near-future hand motion instead of reacting only to the current hand-microrobot distance.'
}

for p in doc.paragraphs:
    text = p.text.strip()
    if text in replacements:
        p.text = replacements[text]

start_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'VI. EXPERIMENTS AND RESULTS':
        start_idx = i
        break
if start_idx is None:
    raise RuntimeError('Could not find VI. EXPERIMENTS AND RESULTS')

body = doc._body._element
start_el = doc.paragraphs[start_idx]._element
children = list(body)
start_child_idx = children.index(start_el)
for child in children[start_child_idx:]:
    if child.tag.endswith('sectPr'):
        continue
    body.remove(child)

def add_heading(text):
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    return p

def add_para(text):
    return doc.add_paragraph(text)

def add_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
    return p

def add_figure(path, caption, width=6.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(caption)

add_heading('VI. EXPERIMENTS AND RESULTS')
add_para('We evaluate the simulation component through two complementary studies. First, a ten-generation league experiment tests whether iterative opponent-pool training improves robustness against heterogeneous learned hand behaviors. Second, a GRU-based network ablation isolates the effect of temporal interaction-history encoding and the auxiliary future-dynamics head. These experiments replace the earlier LSTM ablation and OOD tables in the previous manuscript version, because the current paper scope centers on the no-opponent-ID league result and the final GRU ablation trained against H1-H10 plus scripted hands.')

add_heading('A. League Training Robustness')
add_para('The league experiment uses a ZPD band of 3.5-5.5 workspace units and trains ten robot generations without providing an explicit opponent identity. Each robot generation is evaluated against the learned hand pool H1-H10 and a scripted hand. Time-in-zone (TIS), defined as the fraction of interaction time spent inside the ZPD band, is used as the main robustness metric. We report both mean TIS across learned hands and worst-case TIS to measure the lower tail of performance.')
add_para('Fig. 1 summarizes the final ten-generation league result. The final robot R10 achieves the strongest learned-hand robustness among the evaluated generations. Mean TIS against learned hands increases from 0.188 for R1 to 0.282 for R10, an absolute gain of 0.094. Worst-case learned-hand TIS increases from 0.138 to 0.204, an absolute gain of 0.065. These gains indicate that iterative opponent accumulation improves not only average performance but also robustness against harder learned hand policies.')
add_para('The sorted empirical matrix further supports this interpretation. When robots are ordered from weak to strong and hands from easy to hard, later robot generations occupy higher-TIS regions across a broader set of opponents. The scripted-hand curve remains useful as a sanity check, but the learned-hand pool exposes robustness differences that a single scripted hand would not reveal. This result motivates league-trained virtual patients as a stress-test distribution for adaptive rehabilitation policies.')
add_figure(league_fig, 'Fig. 1. League-training overview for the no-opponent-ID ZPD 3.5-5.5 simulation. The final robot generation improves mean and worst-case time-in-zone against the learned-hand pool, while the sorted empirical matrix visualizes robustness across increasingly difficult virtual patients.')

add_heading('B. GRU Network Ablation')
add_para('The architecture ablation compares three robot policy encoders trained for one generation against the same H1-H10 plus scripted-hand pool. MLP uses only the scalar spatial state and ignores interaction history. GRU adds a temporal sequence encoder over the 16-frame relative-displacement buffer. GRU+Aux further adds the auxiliary future-prediction objective described in Section IV-B. Each group is trained for 8M environment steps. Raw per-episode reward, episode length, ZPD steps, and workspace coverage are recorded independently of TensorBoard.')
add_para('Fig. 2 shows the updated ablation result. The top row uses raw per-episode training data smoothed over a 200k-timestep window. The MLP baseline plateaus at lower episode reward and shorter episode length than both recurrent variants. In the final 10% of training episodes, MLP obtains a mean reward of -42.94 and mean episode length of 27.58, whereas GRU reaches -35.56 reward and 36.12 episode length. This corresponds to an improvement of 7.39 reward points and 8.54 steps per episode. GRU+Aux reaches a similar final-stage reward (-35.39) and episode length (35.92), indicating that temporal sequence modeling is the dominant contributor to aggregate control performance in this run.')
add_para('The auxiliary model does not clearly dominate GRU-only on final aggregate metrics, but it provides an interpretable representation-learning signal. The bottom row of Fig. 2 visualizes future-motion prediction examples and the trajectory error over the eight-step horizon. The error increases with prediction horizon, as expected for multi-step forecasting, while early-step predictions remain close to the observed hand trajectory. These results support the auxiliary task as a mechanism for encouraging the encoder to internalize latent patient dynamics, even though additional tuning may be required for a larger control-performance gain.')
add_figure(ablation_fig, 'Fig. 2. GRU ablation and auxiliary future-motion prediction. Panels (a) and (b) show smoothed raw training reward and episode length for MLP, GRU, and GRU+Aux policies trained against the H1-H10 plus scripted-hand pool. Panels (c) and (d) visualize representative auxiliary future-prediction examples and horizon-dependent trajectory error.')

add_heading('TABLE II: FINAL-STAGE GRU ABLATION SUMMARY')
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
for cell, value in zip(table.rows[0].cells, ['Method', 'Reward', 'Episode Length', 'ZPD Steps', 'Workspace Coverage']):
    cell.text = value
for row in [
    ['MLP', '-42.94', '27.58', '18.47', '0.297'],
    ['GRU', '-35.56', '36.12', '22.08', '0.347'],
    ['GRU+Aux', '-35.39', '35.92', '21.96', '0.343'],
]:
    cells = table.add_row().cells
    for cell, value in zip(cells, row):
        cell.text = value
add_para('Table II reports the final-stage statistics from the last 10% of training episodes. Both recurrent variants produce higher ZPD steps and workspace coverage than the MLP baseline. GRU-only obtains 22.08 ZPD steps and 0.347 workspace coverage, while GRU+Aux obtains 21.96 ZPD steps and 0.343 workspace coverage. The close GRU and GRU+Aux values suggest that the current auxiliary objective primarily acts as representation supervision rather than producing a clear aggregate-performance advantage over GRU-only training.')

add_heading('C. Sim-to-Real Transfer')
add_para('The physical UR10 deployment is not directly comparable to simulation through a pointwise sim-real gap table, because real human behavior, camera latency, magnetic actuation, and safety constraints are not distribution-matched to virtual hands. We therefore treat sim-to-real validation as a real-world closed-loop feasibility study rather than as a paired simulation benchmark. The implementation uses the same 44-dimensional observation structure as simulation: robot position from RTDE, hand and microrobot positions from the overhead camera, boundary features, previous action, and the displacement-history buffer.')
add_para('The implementation uses temporal decoupling between perception and control. The camera runs at 10-15 Hz, but the robot needs deterministic commands at a fixed control rate. Running both in one thread would cause jitter: a slow YOLO inference would stall the control loop and freeze the robot. We therefore use two threads connected by a single-element message queue. The vision thread captures frames, runs undistortion and detection, and pushes results to the queue. If the queue is full, the stale result is replaced. The control thread runs at 20 Hz and attempts a non-blocking read each cycle. If a new frame is available, the hand position is updated; otherwise the system falls back to dead reckoning using a low-pass filtered velocity estimate from recent vision updates.')
add_para('The 44-dimensional observation is assembled from the estimated hand position, the robot position via RTDE, and the displacement-history buffer, then fed to the PPO policy. The 2D action is accumulated as a virtual target in pixel space, rate-limited to prevent large per-step motion, and mapped to world coordinates via the inverse homography. The final pose is sent to the UR10 through the servoL interface with a 0.1 s lookahead time and proportional gain of 400, yielding smooth end-effector tracking.')

add_heading('VII. DISCUSSION')
add_para('The revised simulation results support two main conclusions. First, league training improves robustness by exposing the robot to a distribution of learned patient behaviors rather than a single scripted controller. The improvement in both mean and worst-case TIS suggests that opponent-pool training reduces exploitability and better approximates patient variability. Second, temporal interaction history is essential for adaptive control in this task. The gap between MLP and GRU policies shows that instantaneous spatial state alone is insufficient for inferring patient motion trends and maintaining sustained interaction.')
add_para('The auxiliary future-prediction task provides a mechanistic and visually interpretable way to verify that the temporal encoder is learning patient-motion dynamics. However, the current GRU+Aux run does not substantially exceed the GRU-only model in final reward, episode length, or ZPD steps. This indicates that auxiliary supervision is useful but not yet fully optimized. Future work should tune the auxiliary loss weights, prediction horizon, and training curriculum to better convert representation quality into control-performance improvement.')

add_heading('VIII. CONCLUSION')
add_para('This manuscript presents a simulation-to-real RL framework for adaptive non-contact upper-limb rehabilitation. CMD-DR generates diverse virtual patient behavior by separating intended movement strategy from motor execution constraints. A dual-stream recurrent encoder uses interaction history to infer latent patient dynamics, and an auxiliary future-prediction head supplies dense self-supervised representation learning. Iterative league training improves robustness across learned hand opponents, and the GRU ablation confirms the importance of temporal sequence modeling. The next stage is to integrate these simulation results with systematic physical validation on the UR10 magnetic microrobot platform.')

doc.save(out)
print(out.as_posix())
