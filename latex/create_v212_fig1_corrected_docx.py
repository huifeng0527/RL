from pathlib import Path
from shutil import copy2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

src = Path("manuscripts/methodology_ieee_v2.11_league_clarity.docx")
out = Path("manuscripts/methodology_ieee_v2.12_fig1_corrected.docx")
league_fig = Path("manuscripts/figures/paper_ready/fig_sim01_league_overview_no_title.png")
ablation_fig = Path("manuscripts/figures/paper_ready/fig_sim02_ablation_gru_aux_composite_filled_no_title.png")
copy2(src, out)
doc = Document(out)

body = doc._body._element
children = list(body)
start_idx = None
for idx, child in enumerate(children):
    if child.tag.endswith("p"):
        text = "".join(node.text for node in child.iter() if node.text).strip()
        if text.startswith("VI. EXPERIMENTS"):
            start_idx = idx
            break
if start_idx is None:
    raise RuntimeError("Could not find Section VI")
for child in children[start_idx:]:
    if child.tag.endswith("sectPr"):
        continue
    body.remove(child)


def add_heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


def add_para(text):
    return doc.add_paragraph(text)


def add_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
    return p


def add_figure(path, caption, width=6.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(caption)


add_heading("VI. EXPERIMENTS AND RESULTS")
add_para("The experiments evaluate the proposed framework from complementary perspectives rather than following the methodology section mechanically. We first define the common simulation setup, training protocols, and evaluation metrics. The main evaluation then examines whether league training improves robustness across learned hand behaviors and how the opponent-sampling process evolves during training. We next analyze the role of temporal interaction history and the auxiliary future-dynamics head through a network ablation. The section closes with the real-world deployment pipeline that connects the simulation policy to the UR10 platform.")

add_heading("A. Experimental Setup")
add_para("All simulation experiments use the pursuit-evasion rehabilitation environment described in Section III. The robot controls the magnetic microrobot, and the hand is controlled either by a scripted pursuit controller or by a learned hand policy. The therapeutic objective is defined by a ZPD distance band of 3.5-5.5 workspace units. A step is counted as in-zone when the hand-microrobot distance lies inside this band.")
add_para("The primary evaluation metric is Time-In-Zone (TIS), the fraction of the episode horizon during which the interaction remains inside the ZPD band. We also report ZPD coverage, episode length, catch rate, and empirical robustness across the learned hand pool. These metrics capture different aspects of the same rehabilitation objective: TIS measures sustained therapeutic interaction over the full horizon, ZPD coverage measures the quality of the realized portion of an episode, episode length captures survival, and catch rate reflects overly easy interactions in which the microrobot is caught too quickly.")
add_para("We compare three representative robot training protocols. The scripted-only baseline is trained against the stochastic scripted pursuit controller. The single-agent baseline is trained against one learned hand policy. The league policy is trained through iterative opponent-pool expansion, where robot generations are trained against a mixture of scripted and learned hand behaviors. For the network ablation, we compare an MLP policy, a GRU policy with temporal interaction history, and a GRU+Aux policy with the auxiliary future-dynamics head.")

add_heading("B. League Training and Robustness Evaluation")
add_para("This evaluation tests whether opponent-pool training reduces the brittle specialization that can occur when the robot is optimized against a single hand model. The league policy is trained for ten generations while the hand pool is expanded with learned hand policies. During robot training, the opponent identity is not given as an input; the robot must infer the current hand behavior from the spatial state and recent displacement history. This design makes the evaluation closer to rehabilitation use, where patient capability is observed through movement rather than provided as a label.")
add_para("Fig. 1 summarizes the league result from both evaluation and training-process perspectives. Panel (a) gives the cross-iteration validation matrix: each row is a robot generation, each column is a learned hand generation, and each cell reports TIS for that robot-hand pairing. This panel shows how performance changes across the full interaction matrix rather than only for the final policy. The stronger values in the first hand column indicate that early learned hands remain easier to regulate, while later columns reveal the harder cases that drive the need for broader opponent coverage.")
add_para("Panel (b) reduces the same validation matrix into generation-level robustness curves. The mean TIS measures average performance across learned hands, the worst-hand TIS identifies the most difficult opponent in the pool, and CVaR20 summarizes the lower tail rather than a single extreme case. The upward trend toward the final generation indicates that league training improves not only the average interaction quality but also the less favorable part of the opponent distribution.")
add_para("Panel (c) presents the robustness frontier by plotting each robot generation in the space of mean TIS and worst-hand TIS, with color indicating CVaR20. This view separates policies that improve average behavior from those that also improve difficult-case behavior. In the rehabilitation setting, this distinction matters because a controller that performs well only for easy virtual patients may still fail to maintain useful interaction for patients with more challenging movement patterns.")
add_para("Panel (d) decomposes the final robot generation across test hands into too-close rate, too-far rate, and ZPD coverage. The stacked bars distinguish two failure modes: interactions that become too easy because the hand catches the microrobot, and interactions that become too difficult because the microrobot remains too far away. The ZPD-coverage curve shows that the final league policy keeps a meaningful portion of each episode in the therapeutic band even when the type of failure varies across hand generations.")
add_para("Panel (e) shows snapshots of the PFSP sampling probabilities during selected training iterations. The curves indicate that opponent sampling is not uniform throughout training; probability mass shifts among learned hands as their relative difficulty changes. This supports the training mechanism behind the robustness result: the robot is repeatedly exposed to opponents that remain informative, while the probability floor prevents older opponents from disappearing entirely from the training distribution.")
add_figure(league_fig, "Fig. 1. League-training validation and opponent-sampling dynamics for the no-opponent-ID ZPD 3.5-5.5 simulation. (a) Cross-iteration TIS matrix between robot generations and learned hand generations. (b) Mean, worst-hand, and CVaR20 TIS across robot generations. (c) Robustness frontier relating mean TIS, worst-hand TIS, and CVaR20. (d) Final-generation failure decomposition showing too-close rate, too-far rate, and ZPD coverage across test hands. (e) PFSP sampling-probability snapshots during selected training iterations.")
add_para("We further test whether the learned robustness transfers across hand-controller types. For this comparison, the final scripted-only, single-agent, and league-trained robot policies are each evaluated against two hand mechanisms: the stochastic scripted pursuit controller and a learned hand agent. The table entries report TIS, ZPD coverage, episode length, and catch rate in that order. Mouse-controlled human-in-the-loop testing is kept separate because it introduces human reaction time and voluntary strategy rather than another automated controller.")
add_heading("TABLE II: POLICY COMPARISON ON SCRIPTED AND LEARNED-AGENT HANDS")
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
for cell, value in zip(table.rows[0].cells, ["Robot policy", "Scripted hand", "Agent hand"]):
    cell.text = value
for row in [
    ["Scripted-only", "0.31 / 0.54 / 52 / 60%", "0.19 / 0.54 / 31 / 90%"],
    ["Single-agent", "0.11 / 0.41 / 25 / 95%", "0.49 / 0.60 / 70 / 38%"],
    ["League", "0.38 / 0.55 / 67 / 43%", "0.48 / 0.65 / 64 / 48%"],
]:
    cells = table.add_row().cells
    for cell, value in zip(cells, row):
        cell.text = value
add_para("Table II shows complementary failure modes for the two non-league baselines. The scripted-only robot remains competitive on the scripted controller but is caught frequently by the learned hand agent, indicating that rule-based pursuit during training does not cover learned strategic behavior. The single-agent robot shows the reverse tendency: it performs well against the learned hand but degrades on the scripted controller. The league-trained robot does not dominate every metric in every column, but it avoids the severe collapse observed in the baselines and provides the most balanced behavior across both automated test conditions. This balance is the relevant outcome for an adaptive rehabilitation controller, whose goal is not to overfit to one hand model but to maintain therapeutic interaction as patient behavior changes.")

add_heading("C. Network Ablation and Auxiliary Dynamics Analysis")
add_para("The ablation study turns from opponent diversity to policy representation. The key question is whether the robot can regulate difficulty from the current geometric state alone, or whether it needs recent patient-motion history. The MLP baseline observes the current geometric state but does not explicitly encode recent interaction history. The GRU policy processes the 16-frame relative-displacement buffer, allowing it to infer hand velocity, response delay, and pursuit tendency. The GRU+Aux policy further adds the future-dynamics prediction head described in Section IV-B.")
add_para("The upper panels of Fig. 2 show that the main performance gain appears when temporal interaction history is introduced. Both recurrent policies learn longer and more rewarding interactions than the MLP baseline. This suggests that the robot benefits from observing how the hand has been moving, not only where it is at the current instant. In this setting, the GRU-only and GRU+Aux policies follow similar aggregate learning trends, indicating that temporal sequence modeling is the dominant contributor to control performance.")
add_para("The lower panels of Fig. 2 provide a complementary interpretation of the auxiliary head. Rather than treating the auxiliary task as a separate source of reward improvement, we use it to inspect whether the temporal encoder has learned predictive structure in hand motion. The predicted trajectories capture short-horizon movement direction and diverge gradually at longer horizons, consistent with the uncertainty of multi-step prediction in an interactive task. Thus, the auxiliary head mainly serves to encourage and expose a temporal representation of patient dynamics.")
add_figure(ablation_fig, "Fig. 2. GRU ablation and auxiliary future-motion prediction. Panels (a) and (b) show smoothed training reward and episode length for MLP, GRU, and GRU+Aux policies trained against the mixed learned-hand and scripted-hand pool. Panels (c) and (d) visualize representative auxiliary future-prediction examples and horizon-dependent trajectory error.")

add_heading("D. Real-World Deployment Pipeline")
add_para("The final part of the experiment section describes how the simulation policy is connected to the physical UR10 platform. This pipeline is presented as an implementation bridge rather than as a paired sim-to-real benchmark, because real hand kinematics, camera latency, magnetic actuation, and safety constraints are not fully represented in the virtual training environment.")
add_para("The physical implementation uses the same 44-dimensional observation structure as simulation: robot position from Real-Time Data Exchange (RTDE), hand and microrobot positions from the overhead camera, boundary features, previous action, and the displacement-history buffer. Perception and control run in separate threads to avoid jitter. The camera runs at 10-15 Hz, whereas the robot requires deterministic commands at a fixed control rate. If both processes ran in one thread, a slow You Only Look Once (YOLO) inference step could stall the control loop.")
add_para("The system therefore uses two threads connected by a single-element message queue. The vision thread captures frames, runs undistortion and detection, and pushes results to the queue. If the queue is full, stale results are replaced. The control thread runs at 20 Hz and attempts a non-blocking read each cycle. If a new frame is available, the hand position is updated; otherwise, the system uses dead reckoning based on a low-pass filtered velocity estimate from recent vision updates. The resulting observation is passed to the PPO policy, and the two-dimensional action is mapped through the inverse homography before being sent to the UR10 through the servoL interface.")


def set_run_font(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        rfonts.set(ns + key, "Times New Roman")


for style in doc.styles:
    if hasattr(style, "font"):
        style.font.name = "Times New Roman"
        style.font.size = Pt(10)

for p in doc.paragraphs:
    for run in p.runs:
        set_run_font(run)

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run)

for section in doc.sections:
    for part in [section.header, section.footer]:
        for p in part.paragraphs:
            for run in p.runs:
                set_run_font(run)

doc.save(out)
print(out.as_posix())
